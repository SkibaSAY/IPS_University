#pip install psutil
#pip install docx
#pip install python-docx

import psutil
import GPUtil
import time
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def get_cpu_usage():
    return psutil.cpu_percent(interval=1)

def get_ram_usage():
    return psutil.virtual_memory().percent

def get_gpu_usage():
    gpus = GPUtil.getGPUs()
    if gpus:
        return gpus[0].load * 100  
    return 0  

def get_processes_by_cpu():
  processes = []
  for proc in psutil.process_iter(['pid', 'name', 'cpu_percent']):
        print(proc.info)
        if proc.info['cpu_percent'] > 0:
            processes.append(proc.info)
  return sorted(processes, key=lambda x: x['cpu_percent'], reverse=True)

def create_report(data):
    document = Document()

    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)


    heading = document.add_paragraph()
    heading_run = heading.add_run("Отчёт о нагрузке системы")
    heading_run.font.size = Pt(16)
    heading_run.bold = True
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph(f"Дата и время создания отчета: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}")
    document.add_paragraph(f"")  

    for entry in data:
        paragraph = document.add_paragraph()
        time_str = entry['time'].strftime('%H:%M:%S')
        text_run = paragraph.add_run(f"Время: {time_str} ")
        text_run.bold = True
        paragraph.add_run(f" | CPU: {entry['cpu']:.2f}% | RAM: {entry['ram']:.2f}% | GPU: {entry['gpu']:.2f}% ")
        
        if entry['cpu'] > 80 or entry['ram'] > 80 or entry['gpu'] > 80:
            text_run.font.highlight_color = 15 # Желтый цвет
            
            process_heading = document.add_paragraph()
            process_heading_run = process_heading.add_run("Процессы с наибольшей нагрузкой:")
            process_heading_run.italic = True
            
            process_paragraph = document.add_paragraph()
            sorted_processes = get_processes_by_cpu()
            for process in sorted_processes[:5]:
                process_paragraph.add_run(f"Имя: {process['name']}, PID: {process['pid']}, CPU: {process['cpu_percent']:.2f}%  | ")
            document.add_paragraph(f"")  

    file_name = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    document.save(file_name)
    print(f"Отчет сохранен в файле: {file_name}")


# def


def main():
    print("Запуск мониторинга...")
    report_data = []
    start_time = time.time()

    while time.time() - start_time < 60:
        cpu_usage = get_cpu_usage()
        ram_usage = get_ram_usage()
        gpu_usage = get_gpu_usage()
        
        report_data.append({
            'time': datetime.now(),
            'cpu': cpu_usage,
            'ram': ram_usage,
            'gpu': gpu_usage
        })
        time.sleep(5)

    #get_processes_by_cpu()
    #print(get_ram_usage())

    #create_report(report_data)
    print(report_data)
    print("Мониторинг завершен. Отчет создан.")

if __name__ == "__main__":
    main()